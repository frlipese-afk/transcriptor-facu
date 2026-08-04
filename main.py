from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, File, UploadFile, Form, Request, Body, BackgroundTasks, Cookie, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
from groq import Groq
from supabase import create_client, Client
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from docx import Document
from io import BytesIO
import boto3
import os
import uuid
import subprocess
import glob
import imageio_ffmpeg
from datetime import datetime, timezone, timedelta

ADMIN_PASSWORD = "narizdequeso" # ¡Cambiá esto por tu contraseña!

app = FastAPI()
# Le decimos a FastAPI dónde están los HTML
templates = Jinja2Templates(directory="templates")

# --- CONFIGURACIÓN ---
client_groq = Groq()
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

s3_client = boto3.client('s3',
    endpoint_url=os.environ.get("R2_ENDPOINT_URL"),
    aws_access_key_id=os.environ.get("R2_ACCESS_KEY"),
    aws_secret_access_key=os.environ.get("R2_SECRET_KEY"),
    region_name='auto'
)
BUCKET_NAME = 'transcriptor-audios'
R2_PUBLIC_URL = os.environ.get("R2_PUBLIC_URL")


# --- PÁGINA PRINCIPAL ---
@app.get("/", response_class=HTMLResponse)
async def leer_inicio(request: Request):
    # Solo le pasamos las llaves de Supabase al HTML para que el login funcione
    return templates.TemplateResponse(request, "index.html", {
        "supabase_url": SUPABASE_URL,
        "supabase_key": SUPABASE_KEY
    })


# Diccionario para guardar el estado de cada transcripción (en memoria)
ESTADOS_TRANSCRIPCION = {}

# --- 1. RUTA QUE RECIBE EL AUDIO Y EMPIEZA A TRABAJAR POR DETRÁS ---
@app.post("/iniciar-transcripcion")
async def iniciar_transcripcion(background_tasks: BackgroundTasks, email: str = Form(...), motor: str = Form(...), audio: UploadFile = File(...)):
    # 1. VERIFICAR Y DESCONTAR CRÉDITOS ANTES DE EMPEZAR
    response_db = supabase.table("usuarios").select("creditos").eq("email", email).execute()
    if not response_db.data:
        raise HTTPException(status_code=403, detail="Tu email no está registrado.")
    
    creditos_actuales = response_db.data[0]['creditos']
    costo = 2 if motor == "groq_voces" else 1
    
    if creditos_actuales < costo:
        raise HTTPException(status_code=402, detail=f"No tenés créditos suficientes. Necesitás {costo} y tenés {creditos_actuales}.")
    
    # Descontamos los créditos ya!
    nuevos_creditos = creditos_actuales - costo
    supabase.table("usuarios").update({"creditos": nuevos_creditos}).eq("email", email).execute()
    supabase.table("usuarios").update({"ultima_actividad": "now()"}).eq("email", email).execute()
    
    # 2. Recién ahora empezamos a trabajar
    job_id = str(uuid.uuid4())
    audio_bytes = await audio.read()
    ESTADOS_TRANSCRIPCION[job_id] = {"estado": "iniciando", "porcentaje": 5}
    
    # Le pasamos los nuevos_creditos a la función de fondo para que sepa cuánto mostrar al final
    background_tasks.add_task(transcribir_en_fondo, job_id, email, audio.filename, audio.content_type, audio_bytes, motor, nuevos_creditos)
    
    return {"job_id": job_id}

# --- 2. LA FUNCIÓN QUE HACE EL TRABAJO DURO (POR DETRÁS) ---
def transcribir_en_fondo(job_id: str, email: str, filename: str, content_type: str, audio_bytes: bytes, motor: str, nuevos_creditos: int):
    costo_creditos = 2 if motor == "groq_voces" else 1
    try:
        # 1. Subir a Cloudflare R2 PRIMERO
        ct = content_type if content_type and content_type.startswith('audio') else 'audio/mpeg'
        nombre_archivo_nube = f"{uuid.uuid4()}_{filename.replace(' ', '_')}"
        s3_client.put_object(Bucket=BUCKET_NAME, Key=nombre_archivo_nube, Body=audio_bytes, ContentType=ct)
        url_audio = f"{R2_PUBLIC_URL}/{nombre_archivo_nube}"
        
        texto_plano = ""
        texto_html = ""

        # 2. TRANSCRIPCIÓN CON GROQ (Para todos los motores)
        extension = filename.split('.')[-1].lower()
        temp_original = f"temp_{job_id}.{extension}"
        with open(temp_original, "wb") as f: f.write(audio_bytes)
        
        ESTADOS_TRANSCRIPCION[job_id] = {"estado": "cortando", "porcentaje": 10}
        
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        subprocess.run([
            ffmpeg_exe, "-i", temp_original, 
            "-f", "segment", "-segment_time", "600", 
            "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le", 
            f"temp_chunk_{job_id}_%03d.wav"
        ], check=True, capture_output=True)
        
        if os.path.exists(temp_original): os.remove(temp_original)
        
        chunks = sorted(glob.glob(f"temp_chunk_{job_id}_*.wav"))
        total_chunks = len(chunks)
        
        modelo_groq = "whisper-large-v3-turbo" if motor == "groq_turbo" else "whisper-large-v3"
        
        for i, chunk_filename in enumerate(chunks):
            ESTADOS_TRANSCRIPCION[job_id] = {
                "estado": "transcribiendo", 
                "actual": i + 1, 
                "total": total_chunks,
                "porcentaje": 10 + int(((i + 1) / total_chunks) * 70)
            }
            
            with open(chunk_filename, "rb") as audio_file:
                response_groq = client_groq.audio.transcriptions.create(
                    model=modelo_groq, 
                    file=audio_file, 
                    response_format="verbose_json"
                )
            
            offset_segundos = i * 600.0
            for segmento in response_groq.segments:
                inicio_real = segmento['start'] + offset_segundos
                minutos = int(inicio_real // 60)
                segundos = int(inicio_real % 60)
                texto_plano += f"[{minutos:02d}:{segundos:02d}] {segmento['text']}\n"
                texto_html += f"<span class='minuto' onclick='saltarA({inicio_real})'>[{minutos:02d}:{segundos:02d}]</span> {segmento['text']}<br>"
            
            os.remove(chunk_filename)
        
        ESTADOS_TRANSCRIPCION[job_id] = {"estado": "guardando", "porcentaje": 90}

        # 3. GUARDADO EN BASE DE DATOS
        titulo_limpio = os.path.splitext(filename)[0]
        supabase.table("transcripciones").insert({
            "user_email": email, 
            "titulo": titulo_limpio, 
            "texto": texto_plano, 
            "audio_url": url_audio,
            "motor_usado": motor
        }).execute()
        
        ESTADOS_TRANSCRIPCION[job_id] = {
            "estado": "terminado", 
            "porcentaje": 100,
            "resultado": {
                "url_audio": url_audio, "texto_html": texto_html, 
                "nuevos_creditos": nuevos_creditos, "email": email
            }
        }
    except Exception as e:
        # Si algo falla a mitad de camino, LE DEVOLVEMOS EL CRÉDITO AL USUARIO
        supabase.table("usuarios").update({"creditos": nuevos_creditos + costo_creditos}).eq("email", email).execute()
        ESTADOS_TRANSCRIPCION[job_id] = {"estado": "error", "mensaje": str(e)}

# --- 3. RUTA QUE LE DICE A LA PÁGINA CÓMO VA ---
@app.get("/estado-transcripcion/{job_id}")
async def estado_transcripcion(job_id: str):
    return ESTADOS_TRANSCRIPCION.get(job_id, {"estado": "no_encontrado"})

# --- 4. RUTA QUE MUESTRA EL RESULTADO FINAL ---
@app.get("/resultado/{job_id}", response_class=HTMLResponse)
async def resultado_transcripcion(job_id: str):
    estado = ESTADOS_TRANSCRIPCION.get(job_id)
    if not estado or estado.get("estado") != "terminado":
        return "<h3>Procesando o error...</h3>"
    
    data = estado["resultado"]
    url_audio = data["url_audio"]
    texto_html = data["texto_html"]
    nuevos_creditos = data["nuevos_creditos"]
    email = data["email"]

    return f"""
    <html>
        <head>
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Resultado</title>
            <style>
                body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; background-color: #f0f2f5; margin: 0; padding: 20px; color: #1a1a1a; }}
                .container {{ max-width: 600px; margin: 20px auto; }}
                .card {{ background: white; border-radius: 16px; padding: 24px; box-shadow: 0 4px 15px rgba(0,0,0,0.08); }}
                h3 {{ margin-top: 0; color: #333; font-size: 22px; }}
                audio {{ width: 100%; margin-bottom: 20px; border-radius: 12px; }}
                .texto {{ white-space: pre-wrap; font-size: 17px; line-height: 1.8; color: #444; }}
                .creditos {{ color: #666; font-size: 14px; margin-top: 20px; border-top: 1px solid #eee; padding-top: 15px; text-align: center; }}
                .btn {{ display: block; width: 100%; padding: 16px; font-size: 17px; border: none; border-radius: 12px; cursor: pointer; font-weight: 600; margin-top: 15px; text-decoration: none; box-sizing: border-box; text-align: center; }}
                .btn-success {{ background: #28a745; color: white; }}
                .btn-back {{ background: #e9ecef; color: #666; }}
                .minuto {{ color: #007bff; cursor: pointer; font-weight: bold; background: #e7f1ff; padding: 3px 8px; border-radius: 6px; margin-right: 5px; display: inline-block; margin-bottom: 5px; font-size: 14px; }}
                .minuto:hover {{ background: #007bff; color: white; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="card">
                    <h3>📝 Transcripción:</h3>
                    <audio id="reproductor" controls src="{url_audio}"></audio>
                    <div class="texto">{texto_html}</div>
                    <div class="creditos">Te quedan <b>{nuevos_creditos}</b> créditos.</div>
                    <a href="/historial?email={email}" class="btn btn-success">📚 Ver mi historial</a>
                    <a href="/" class="btn btn-back">⬅️ Volver a transcribir</a>
                </div>
            </div>
            <script>
                function saltarA(segundos) {{
                    const audio = document.getElementById('reproductor');
                    audio.currentTime = segundos;
                    audio.play();
                }}
            </script>
        </body>
    </html>
    """


# --- HISTORIAL ---
@app.get("/historial", response_class=HTMLResponse)
async def ver_historial(request: Request, email: str):
        # Actualizar última actividad
    supabase.table("usuarios").update({"ultima_actividad": "now()"}).eq("email", email).execute()
    
    response = supabase.table("transcripciones").select("id, titulo, texto, fecha, audio_url").eq("user_email", email).order("fecha", desc=True).execute()
    notas = response.data
    
    notas_procesadas = []
    for nota in notas:
        texto_html = ""
        # Usamos .get() por si la nota no tiene texto (evita el error)
        texto_original = nota.get('texto', '') or ''
        
        for linea in texto_original.split('\n'):
            # Si la línea tiene corchete, intentamos sacarle el minuto
            if linea.startswith('[') and ']' in linea:
                partes = linea.split(']', 1)
                minuto_str = partes[0].replace('[', '').strip()
                try:
                    # Intentamos convertir el minuto a números
                    mins, segs = map(int, minuto_str.split(':'))
                    segundos_totales = mins * 60 + segs
                    texto_html += f"<span class='minuto' onclick='saltarAHistorial({segundos_totales}, {nota['id']})'>[{minuto_str}]</span>{partes[1]}<br>"
                except:
                    # Si falla (ej: no son números), mostramos la línea normal
                    texto_html += f"{linea}<br>"
            else:
                # Si no tiene corchete, la mostramos normal
                texto_html += f"{linea}<br>"
        
        nota['texto_html'] = texto_html
        notas_procesadas.append(nota)

    return templates.TemplateResponse(request, "historial.html", {
        "notas": notas_procesadas,
        "email": email
    })


# --- PANEL DE ADMINISTRACIÓN SECRETO (CON CONTRASEÑA) ---
@app.get("/admin-secreto", response_class=HTMLResponse)
async def panel_admin(request: Request, admin_auth: str = Cookie(None)):
    # 1. Revisar si tiene la cookie de admin
    if admin_auth != "autorizado":
        # Si no la tiene, mostrar formulario de contraseña
        return """
        <html><head><meta name="viewport" content="width=device-width, initial-scale=1.0"><title>Admin Login</title>
        <style>body { font-family: Arial; background: #e9ecef; display: flex; justify-content: center; align-items: center; height: 100vh; margin: 0; } .card { background: white; padding: 40px; border-radius: 16px; box-shadow: 0 4px 15px rgba(0,0,0,0.1); text-align: center; } input { padding: 15px; width: 100%; box-sizing: border-box; margin: 10px 0; border-radius: 8px; border: 1px solid #ccc; } button { background: #2563eb; color: white; border: none; padding: 15px; width: 100%; border-radius: 8px; cursor: pointer; font-weight: bold; }</style></head>
        <body>
            <div class="card">
                <h2>🔒 Acceso Restringido</h2>
                <form action="/admin-login" method="post">
                    <input type="password" name="passw" placeholder="Contraseña de Admin" required>
                    <button type="submit">Entrar al panel</button>
                </form>
            </div>
        </body></html>
        """
    
    # 2. Si tiene la cookie, mostrar el panel real
    response = supabase.table("usuarios").select("*").execute()
    usuarios = response.data
    ahora = datetime.now(timezone.utc)
    
    tabla_html = ""
    for u in usuarios:
        ultima_str = "Nunca"
        estado_dot = "⚪"
        if u.get('ultima_actividad'):
            try:
                ultima_dt = datetime.fromisoformat(u['ultima_actividad'].replace('Z', '+00:00'))
                diff = ahora - ultima_dt
                if diff < timedelta(minutes=5): estado_dot = "🟢"
                elif diff < timedelta(hours=1): estado_dot = "🟡"
                else: estado_dot = "🔴"
                ultima_str = ultima_dt.astimezone().strftime('%d/%m %H:%M')
            except: pass

        tabla_html += f"""
        <tr>
            <td style='font-size: 14px;'>{estado_dot} {u['email']}</td>
            <td style='text-align: center;'><strong style='font-size: 18px; color: #2563eb;'>{u['creditos']}</strong></td>
            <td style='font-size:12px; color:#666;'>{ultima_str}</td>
            <td style='white-space: nowrap; text-align: center;'>
                <!-- Botones rápidos -->
                <form action='/agregar-creditos' method='post' style='display:inline;'>
                    <input type='hidden' name='email' value='{u['email']}'>
                    <input type='hidden' name='cantidad' value='5'>
                    <button type='submit' style='background:#16a34a; color:white; border:none; border-radius:5px; cursor:pointer; padding:6px 12px; font-weight: bold;'>➕ 5</button>
                </form>
                <form action='/agregar-creditos' method='post' style='display:inline; margin-left: 5px;'>
                    <input type='hidden' name='email' value='{u['email']}'>
                    <input type='hidden' name='cantidad' value='-5'>
                    <button type='submit' style='background:#dc3545; color:white; border:none; border-radius:5px; cursor:pointer; padding:6px 12px; font-weight: bold;'>➖ 5</button>
                </form>
                <br>
                <!-- Input manual -->
                <form action='/agregar-creditos' method='post' style='display:inline-flex; margin-top: 8px; align-items: center; gap: 5px;'>
                    <input type='hidden' name='email' value='{u['email']}'>
                    <input type='number' name='cantidad' value='10' style='width: 55px; padding: 5px; border-radius: 5px; border: 1px solid #ccc; text-align: center;'>
                    <button type='submit' style='background:#2563eb; color:white; border:none; border-radius:5px; cursor:pointer; padding:6px 12px; font-weight: bold;'>✔</button>
                </form>
            </td>
        </tr>"""

    return f"""
    <html><head><meta name="viewport" content="width=device-width, initial-scale=1.0"><title>Admin</title><style>body {{ font-family: Arial; padding: 20px; background: #e9ecef; }} .contenedor {{ max-width: 800px; margin: auto; background: white; padding: 20px; border-radius: 8px; }} table {{ border-collapse: collapse; width: 100%; margin-top: 20px; }} th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }} th {{ background-color: #f2f2f2; }}</style></head>
    <body><div class="contenedor"><h1>Panel de Administración 🛠️</h1><table><tr><th>Usuario</th><th>Créditos</th><th>Última Actividad</th><th>Acciones</th></tr>{tabla_html}</table></div></body></html>
    """

# --- RUTA PARA VALIDAR LA CONTRASEÑA ---
@app.post("/admin-login")
async def admin_login(request: Request):
    # Leemos el formulario de la forma más directa
    form = await request.form()
    passw = form.get("passw")
    
    if passw == ADMIN_PASSWORD:
        resp = RedirectResponse(url="/admin-secreto", status_code=303)
        resp.set_cookie(key="admin_auth", value="autorizado", max_age=86400) # Dura 1 día
        return resp
    
    return "<h3>Contraseña incorrecta. <a href='/admin-secreto'>Volver</a></h3>" 

@app.post("/agregar-creditos")
async def agregar_creditos(email: str = Form(...), cantidad: int = Form(...)):
    res = supabase.table("usuarios").select("creditos").eq("email", email).execute()
    if res.data:
        creditos_actuales = res.data[0]['creditos']
        nuevos = creditos_actuales + cantidad
        supabase.table("usuarios").update({"creditos": nuevos}).eq("email", email).execute()
    else:
        supabase.table("usuarios").insert({"email": email, "creditos": cantidad}).execute()
    
    return "<h3>¡Créditos actualizados! Volvé al panel.</h3><a href='/admin-secreto'>Volver al panel</a>"

    # --- BORRAR TRANSCRIPCIÓN ---
@app.post("/borrar/{transcripcion_id}")
async def borrar_transcripcion(transcripcion_id: int, email: str = Form(...)):
    # 1. Buscamos la nota en Supabase para conseguir el link del audio
    response = supabase.table("transcripciones").select("audio_url").eq("id", transcripcion_id).execute()
    
    if response.data:
        audio_url = response.data[0].get('audio_url', '')
        
        # 2. Si tiene audio, lo borramos de Cloudflare R2
        if audio_url:
            # Agarramos el nombre del archivo (lo que está después de la última barra /)
            nombre_archivo = audio_url.split('/')[-1]
            try:
                s3_client.delete_object(Bucket=BUCKET_NAME, Key=nombre_archivo)
            except Exception as e:
                print(f"Error al borrar audio de R2: {e}")
        
        # 3. Borramos el texto de la base de datos (Supabase)
        supabase.table("transcripciones").delete().eq("id", transcripcion_id).execute()
    
    # 4. Volvemos al historial
    return RedirectResponse(url=f"/historial?email={email}", status_code=303)


# --- EDITAR TÍTULO ---
@app.post("/editar-titulo/{transcripcion_id}")
async def editar_titulo(transcripcion_id: int, data: dict = Body(...)):
    email = data.get("email")
    nuevo_titulo = data.get("titulo")
    nuevo_titulo = os.path.splitext(nuevo_titulo)[0]
    
    # Actualizamos en Supabase
    supabase.table("transcripciones").update({"titulo": nuevo_titulo}).eq("id", transcripcion_id).eq("user_email", email).execute()
    
     # Le respondemos a JavaScript que salió todo bien
    return {"titulo": nuevo_titulo}

    # --- EDITAR TEXTO ---
@app.get("/editar/{transcripcion_id}", response_class=HTMLResponse)
async def editar_nota(request: Request, transcripcion_id: int, email: str):
    response = supabase.table("transcripciones").select("*").eq("id", transcripcion_id).eq("user_email", email).execute()
    if not response.data:
        return "<h3>No se encontró la nota.</h3>"
    
    nota = response.data[0]
    
    texto_html = ""
    texto_original = nota.get('texto', '') or ''
    for linea in texto_original.split('\n'):
        if linea.startswith('[') and ']' in linea:
            partes = linea.split(']', 1)
            minuto_str = partes[0].replace('[', '').strip()
            try:
                mins, segs = map(int, minuto_str.split(':'))
                segundos_totales = mins * 60 + segs
                # Usamos las clases nuevas del diseño: transcript-line, timestamp y editable-text
                texto_html += f"<div class='transcript-line'><span class='timestamp' contenteditable='false' onclick='saltarA({segundos_totales})'>[{minuto_str}]</span><span class='editable-text'>{partes[1]}</span></div>"
            except:
                texto_html += f"<div class='transcript-line'><span class='editable-text'>{linea}</span></div>"
        else:
            texto_html += f"<div class='transcript-line'><span class='editable-text'>{linea}</span></div>"
    
    nota['texto_html'] = texto_html

    return templates.TemplateResponse(request, "editar.html", {
        "nota": nota,
        "email": email
    })
    
    # Preparamos el texto para la caja editable
    texto_html = ""
    texto_original = nota.get('texto', '') or ''
    for linea in texto_original.split('\n'):
        if linea.startswith('[') and ']' in linea:
            partes = linea.split(']', 1)
            minuto_str = partes[0].replace('[', '').strip()
            try:
                mins, segs = map(int, minuto_str.split(':'))
                segundos_totales = mins * 60 + segs
                # Le ponemos contenteditable="false" al minutito para que no se pueda borrar
                texto_html += f"<div class='linea'><span class='minuto' contenteditable='false' onclick='saltarA({segundos_totales})'>[{minuto_str}]</span><span class='texto-editable'>{partes[1]}</span></div>"
            except:
                texto_html += f"<div class='linea'>{linea}</div>"
        else:
            texto_html += f"<div class='linea'>{linea}</div>"
    
    nota['texto_html'] = texto_html

    return templates.TemplateResponse(request, "editar.html", {
        "nota": nota,
        "email": email
    })

@app.post("/guardar-edicion/{transcripcion_id}")
async def guardar_edicion(transcripcion_id: int, email: str = Form(...), titulo: str = Form(...), texto: str = Form(...)):
    # Actualizamos el texto y el título en Supabase
    titulo_limpio = os.path.splitext(titulo)[0] # Le sacamos la extensión por las dudas
    supabase.table("transcripciones").update({
        "titulo": titulo_limpio,
        "texto": texto
    }).eq("id", transcripcion_id).eq("user_email", email).execute()
    
    return RedirectResponse(url=f"/historial?email={email}", status_code=303)


# --- DESCARGAR EN WORD ---
import re 

@app.get("/descargar/{transcripcion_id}")
async def descargar_word(transcripcion_id: int, con_minutos: bool = True):
    response = supabase.table("transcripciones").select("titulo, texto").eq("id", transcripcion_id).execute()
    if not response.data:
        return "No encontrado"
    
    nota = response.data[0]
    titulo = nota['titulo']
    texto = nota['texto']
    
    # Si el usuario eligió "Sin minutos", se los borramos
    if not con_minutos:
        texto = re.sub(r'\[\d{2}:\d{2}\]\s*', '', texto)
    
    doc = Document()
    doc.add_heading(titulo, 0)
    
    for linea in texto.split('\n'):
        doc.add_paragraph(linea)
    
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    headers = {
        'Content-Disposition': f'attachment; filename="{titulo}.docx"'
    }
    return StreamingResponse(file_stream, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)